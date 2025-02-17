from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

def generar_estado_cuenta(datos, numero_apartamento):
    try:
        # Crear el documento
        doc = Document()

        # Título
        titulo = doc.add_heading(level=1)
        titulo_run = titulo.add_run(f"Estado de Cuenta - Apartamento {numero_apartamento}")
        titulo_run.bold = True
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Decoración inicial
        doc.add_paragraph("\n")
        decoracion = doc.add_paragraph()
        decoracion_run = decoracion.add_run("============================================")
        decoracion_run.bold = True
        decoracion.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Espacio
        doc.add_paragraph("\n")

        # Renuncia
        renuncia = doc.add_paragraph()
        renuncia.add_run("Nota Importante:\n").bold = True
        renuncia.add_run(
            "Las fechas indicadas en este estado de cuenta son aproximadas, ya que reflejan la fecha de depósito en la cuenta bancaria. "
            "Este documento es para los usos que el receptor considere convenientes. También se detalla que los pagos debieron ser realizados "
            "el día 12 de cada mes según contrato, pero se realizaron en las fechas indicadas."
        )
        renuncia.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Espacio
        doc.add_paragraph("\n")

        # Tabla
        tabla = doc.add_table(rows=1, cols=5)
        tabla.style = 'Table Grid'

        # Encabezados de la tabla
        encabezados = ['Fecha', 'Destino', 'Apartamento', 'Mes', 'Monto']
        hdr_cells = tabla.rows[0].cells
        for i, encabezado in enumerate(encabezados):
            cell = hdr_cells[i]
            cell.text = encabezado
            run = cell.paragraphs[0].runs[0]
            run.bold = True
            run.font.size = Pt(10)

        # Agregar datos
        for fila in datos:
            if len(fila) != 5:
                print(f"Advertencia: La fila {fila} no tiene 5 elementos.")
                continue
            row_cells = tabla.add_row().cells
            for i, valor in enumerate(fila):
                row_cells[i].text = str(valor)

        # Espacio
        doc.add_paragraph("\n")

        # Decoración final
        decoracion_final = doc.add_paragraph()
        decoracion_final.add_run("============================================").bold = True
        decoracion_final.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Firma
        firma = doc.add_paragraph("\n\n")
        firma.add_run("_______________________________\n").bold = True
        firma.add_run("Firma del Administrador")
        firma.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Guardar el documento
        nombre_archivo = f"Estado_Cuenta_Apartamento_{numero_apartamento}.docx"
        doc.save(nombre_archivo)
        print(f"El documento '{nombre_archivo}' se ha generado correctamente.")

    except Exception as e:
        print(f"Error al generar el documento: {e}")

# Ejemplo de datos
estado_cuenta_datos = [
    ["06/may/2024",	"DEPOSITO APARTAMENTO 3",	3	,"mayo",	"$160.00"],
    ["03/jun/2024",	"PAGO MENSUALIDAD APARTAMENTO 3",	3	,"junio",	"$160.00"],
    ["04/jul/2024",	"PAGO MENSUALIDAD APARTAMENTO 3",	3	,"julio",	"$160.00"],
    ["09/ago/2024",	"PAGO MENSUALIDAD APARTAMENTO 3",	3	,"agosto",	"$160.00"],
    ["10/sep/2024",	"PAGO MENSUALIDAD APARTAMENTO 3",	3	,"septiembre",	"$290.00"],
    ["07/oct/2024",	"PAGO MENSUALIDAD APARTAMENTO 3",	3	,"octubre",	"$160.00"]
]


# Generar el estado de cuenta
generar_estado_cuenta(estado_cuenta_datos, 3)
