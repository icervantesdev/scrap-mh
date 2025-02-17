from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generar_recibos(datos):
    # Solicitar el número del apartamento
    numero_apartamento = int(input("Ingrese el número del apartamento: "))

    # Filtrar los datos por el número del apartamento
    datos_filtrados = [fila for fila in datos if fila[2] == numero_apartamento]

    if not datos_filtrados:
        print(f"No hay datos para el apartamento {numero_apartamento}.")
        return

    # Crear el documento
    doc = Document()

    for fila in datos_filtrados:
        fecha, destino, apartamento, mes, monto = fila

        # Añadir el contenido del recibo
        recibo = doc.add_paragraph()
        recibo_run = recibo.add_run(f"RECIBO DE ARRENDAMIENTO\n")
        recibo_run.bold = True
        recibo_run.font.size = Pt(14)
        recibo.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f"Fecha: {fecha}")
        doc.add_paragraph(
            f"Recibí de arrendatarios del Apartamento {apartamento}: {numero_a_texto(apartamento)}, La cantidad de: {monto_a_texto(monto)} ({monto}) DOLARES AMERICANOS. En concepto de: {destino}."
        )

        # Espacio para la firma
        firma = doc.add_paragraph("\n\n___________________________\nOscar Ivan Cervantes Guevara")
        firma.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Añadir un salto de página después de cada recibo (cada media página)
        doc.add_paragraph("\n\n")

    # Guardar el documento
    nombre_archivo = f"Recibos_Apartamento_{numero_apartamento}.docx"
    doc.save(nombre_archivo)
    print(f"El documento '{nombre_archivo}' se ha generado correctamente.")

# Función para convertir números a texto
# (Esta es una implementación básica y se puede mejorar para manejar números más grandes o complejos)
def numero_a_texto(numero):
    numeros = {
        1: "UNO",
        2: "DOS",
        3: "TRES",
        4: "CUATRO",
        5: "CINCO",
        # Agrega más si es necesario
    }
    return numeros.get(numero, str(numero).upper())

def monto_a_texto(monto):
    monto = monto.replace("$", "").replace(",", "")
    partes = monto.split(".")
    entero = int(partes[0])
    centavos = int(partes[1]) if len(partes) > 1 else 0

    texto = f"{numero_a_texto(entero)} DOLARES"
    if centavos > 0:
        texto += f" CON {numero_a_texto(centavos)} CENTAVOS"
    return texto

# Datos de ejemplo
datos = [
    ["06/may/2024", "PAGO MENSUALIDAD APARTAMENTO 1", 1, "mayo", "$100.00"],
    ["06/may/2024", "DEPOSITO APARTAMENTO 3", 3, "mayo", "$160.00"],
    ["03/jun/2024", "PAGO MENSUALIDAD APARTAMENTO 3", 3, "junio", "$160.00"],
    ["03/jun/2024", "PAGO MENSUALIDAD APARTAMENTO 5", 5, "junio", "$160.00"],
    ["03/jul/2024", "PAGO MENSUALIDAD APARTAMENTO 1", 1, "julio", "$125.00"],
    ["03/jul/2024", "PAGO MENSUALIDAD APARTAMENTO 5", 5, "julio", "$160.00"],
    ["04/jul/2024", "PAGO MENSUALIDAD APARTAMENTO 3", 3, "julio", "$160.00"],
    ["16/jul/2024", "DEPOSITO APARTAMENTO 2", 2, "julio", "$100.00"],
    ["16/jul/2024", "PAGO MENSUALIDAD APARTAMENTO 2", 2, "julio", "$125.00"],
    ["16/jul/2024", "PAGO MENSUALIDAD APARTAMENTO 4", 4, "julio", "$150.00"],
    ["16/jul/2024", "DEPOSITO APARTAMENTO 4", 4, "julio", "$150.00"],
    ["09/ago/2024", "PAGO MENSUALIDAD APARTAMENTO 3", 3, "agosto", "$160.00"],
    ["11/ago/2024", "PAGO MENSUALIDAD APARTAMENTO 1", 1, "agosto", "$125.00"],
    ["15/ago/2024", "PAGO MENSUALIDAD APARTAMENTO 2", 2, "agosto", "$125.00"],
    ["17/ago/2024", "PAGO MENSUALIDAD APARTAMENTO 5", 5, "agosto", "$160.00"],
    ["21/ago/2024", "PAGO MENSUALIDAD APARTAMENTO 4", 4, "agosto", "$150.00"],
    ["10/sep/2024", "PAGO MENSUALIDAD APARTAMENTO 3", 3, "septiembre", "$290.00"],
    ["10/sep/2024", "PAGO MENSUALIDAD APARTAMENTO 5", 5, "septiembre", "$130.00"],
    ["19/sep/2024", "PAGO MENSUALIDAD APARTAMENTO 2", 2, "septiembre", "$125.00"],
    ["07/oct/2024", "PAGO MENSUALIDAD APARTAMENTO 1", 1, "octubre", "$125.00"],
    ["07/oct/2024", "PAGO MENSUALIDAD APARTAMENTO 3", 3, "octubre", "$160.00"],
    ["07/oct/2024", "PAGO MENSUALIDAD APARTAMENTO 5", 5, "octubre", "$125.00"],
    ["31/oct/2024", "PAGO MENSUALIDAD APARTAMENTO 2", 2, "octubre", "$125.00"],
    ["05/nov/2024", "PAGO MENSUALIDAD APARTAMENTO 5", 5, "noviembre", "$100.00"],
    ["08/nov/2024", "PAGO MENSUALIDAD APARTAMENTO 1", 1, "noviembre", "$125.00"],
    ["17/nov/2024", "PAGO MENSUALIDAD APARTAMENTO 5", 5, "noviembre", "$60.00"],
    ["05/dic/2024", "PAGO MENSUALIDAD APARTAMENTO 2", 2, "diciembre", "$125.00"],
    ["18/dic/2024", "PAGO MENSUALIDAD APARTAMENTO 5", 5, "diciembre", "$160.00"],
    ["29/dic/2024", "PAGO MENSUALIDAD APARTAMENTO 1", 1, "diciembre", "$125.00"],
    ["29/dic/2024", "PAGO MENSUALIDAD APARTAMENTO 2", 2, "diciembre", "$20.00"],
    ["29/dic/2024", "PAGO MENSUALIDAD APARTAMENTO 2", 2, "diciembre", "$125.00"]
]

# Generar recibos
generar_recibos(datos)
